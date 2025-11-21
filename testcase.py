import datetime
import json
from google import genai
from config import Config  # api_key của bạn nằm trong file config.py
from docx import Document
from docx.shared import Inches

from googlegenai import test_ask_gemini

# ==================== DANH SÁCH TEST CASE ====================
test_inputs = [
    # "Họp team vào 9h sáng ngày mai tại phòng họp A, nhắc trước 15 phút",
    # "Hội thảo từ 8h đến 17h ngày 25/12/2024 tại Trung tâm Hội nghị",
    # "Sinh nhật bạn An lúc 18h30 đến 21h thứ 7 tuần sau tại nhà hàng",
    # "Meeting với client lúc 14h chiều mai 2 tiếng qua Zoom",
    # "Du lịch từ 7h sáng 01/01/2025 đến 19h tối 03/01/2025",
    # "hop team luc 9h sang mai tai phong hop A",
    # "hen ban An 6h chieu t7 nay tai CF Highlands",
    # "su kien ra mat san pham tu 8h den 12h ngay 15/12 tai Bitexco",
    # "di choi vs crush 7h toi mai tai rap phim CGV",
    # "hop khach hang luc 2h chieu tai van phong",
    # "Họp team T2 tuần sau 9h sáng tại phòng 301",
    # "Phỏng vấn T4 14h30 văn phòng HN",
    # "Đi ăn ngày mốt 12h trưa tại nhà hàng Sushi",
    # "Training tuần sau thứ 3 từ 8h đến 17h",
    # "Đi gym tối nay 19h tại California Fitness",
    # "Họp team",
    # "Meeting sáng mai",
    # "Gặp bạn 3h chiều mai",
    # "Phỏng vấn 10h sáng 20/12 tại văn phòng",
    # "Sinh nhật em ngày 25/12",
    "hẹn ny 7h tối nay xem phim, nhớ nhắc 30p",
    "đi ăn vs crush 6h30 chiều mai tai quan lẩu",
    "mai t đi cafe vs m lúc 2h ok k",
    "Meeting team sale 9am tomorrow tại office",
    "chiều nay đi chill 5h tại Starbucks",
    "Hẹn giữa buổi sáng mai tại quán trà sữa",
    "Workshop lúc 08:00 AM ngày 30-12-2024 tại Trung tâm Đào tạo",
    "Nhắc tôi tham dự lễ trao giải bắt đầu lúc 19h tối thứ 6 tuần này tại Nhà hát Thành phố",
    "Có thể ghi lịch họp lúc 10h sáng thứ 5 giúp tôi được không?",
    "Nhắc tôi phỏng vấn 2h chiều 28/12 tại tầng 15 Times Square, chuẩn bị trước 1 tiếng"
]
# ==================== CHẠY TEST & LƯU KẾT QUẢ ====================
results = []

print("Đang chạy test case...\n")
for i, inp in enumerate(test_inputs, 1):
    print(f"{i}/{len(test_inputs)}: {inp[:60]}...")
    result = test_ask_gemini(inp)
    results.append({
        "stt": i,
        "input": inp,
        "json": json.dumps(result, ensure_ascii=False, indent=2)
    })

# ==================== TẠO FILE WORD ====================
doc = Document()
doc.add_heading('Kết quả Test Trích xuất Lịch từ Tiếng Việt (Gemini 1.5 Flash)', 0)

p = doc.add_paragraph('Thời gian test: ')
p.add_run(datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")).bold = True
doc.add_paragraph(f'Tổng số test case: {len(results)}')
doc.add_paragraph()

# Tạo bảng
table = doc.add_table(rows=1, cols=3, style='Table Grid')
table.autofit = True
table.allow_autofit = True

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'STT'
hdr_cells[1].text = 'Input'
hdr_cells[2].text = 'Kết quả JSON'

# Thêm dữ liệu
for item in results:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item["stt"])
    row_cells[1].text = item["input"]
    row_cells[2].text = item["json"]

# Điều chỉnh độ rộng cột
table.columns[0].width = Inches(0.5)
table.columns[1].width = Inches(4.0)
table.columns[2].width = Inches(5.5)

# Lưu file
output_file = "Ket_qua_test_trich_xuat_lich_tieng_Viet.docx"
doc.save(output_file)

print(f"\nHoàn thành! Đã xuất ra file: {output_file}")